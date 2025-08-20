# Production AI User Story Generator Agent
# Complete implementation with FastAPI, LangChain, RAG, and enterprise features

# ========================================
# 1. BACKEND IMPLEMENTATION - main.py
# ========================================

from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Security
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import asyncio
import logging
import os
import json
import hashlib
from datetime import datetime, timedelta
import uuid

# LangChain and AI imports
from langchain.embeddings import OpenAIEmbeddings, AzureOpenAIEmbeddings
from langchain.llms import OpenAI, AzureOpenAI
from langchain.chat_models import ChatOpenAI, AzureChatOpenAI, ChatGooglePalm
from langchain.vectorstores import Chroma, FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import PDFLoader, TextLoader, UnstructuredWordDocumentLoader
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.schema import Document
from langchain.memory import ConversationBufferMemory
from langchain_experimental.agents import create_pandas_dataframe_agent

# Vector database and storage
import chromadb
from chromadb.config import Settings
import redis
from sqlalchemy import create_engine, Column, String, DateTime, Text, Integer, Boolean
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session

# Document processing
import PyPDF2
import docx
import pandas as pd
import requests
from bs4 import BeautifulSoup

# Security and monitoring
import jwt
from passlib.context import CryptContext
from prometheus_client import Counter, Histogram, generate_latest
import bleach

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ========================================
# 2. CONFIGURATION MODELS
# ========================================

class LLMConfig(BaseModel):
    provider: str = Field(..., description="LLM provider: openai, azure, gemini, anthropic")
    model_name: str = Field(..., description="Model name (gpt-4, gpt-3.5-turbo, etc.)")
    api_key: str = Field(..., description="API key for the provider")
    base_url: Optional[str] = Field(None, description="Base URL for Azure/custom endpoints")
    api_version: Optional[str] = Field(None, description="API version for Azure")
    temperature: float = Field(0.7, ge=0, le=2, description="Model temperature")
    max_tokens: int = Field(2000, ge=1, le=8000, description="Maximum tokens")

class RAGConfig(BaseModel):
    chunk_size: int = Field(1000, ge=100, le=4000)
    chunk_overlap: int = Field(200, ge=0, le=1000)
    vector_store: str = Field("chroma", description="Vector store: chroma, faiss, pinecone")
    similarity_threshold: float = Field(0.7, ge=0, le=1)
    max_retrieved_docs: int = Field(5, ge=1, le=20)

class SecurityConfig(BaseModel):
    jwt_secret: str = Field(..., description="JWT secret key")
    jwt_algorithm: str = Field("HS256")
    access_token_expire_minutes: int = Field(30)
    rate_limit_requests: int = Field(100)
    rate_limit_window: int = Field(3600)  # 1 hour

class IntegrationConfig(BaseModel):
    confluence_base_url: Optional[str] = None
    confluence_username: Optional[str] = None
    confluence_api_token: Optional[str] = None
    jira_base_url: Optional[str] = None
    jira_username: Optional[str] = None
    jira_api_token: Optional[str] = None

class AppConfig(BaseModel):
    llm: LLMConfig
    rag: RAGConfig
    security: SecurityConfig
    integrations: IntegrationConfig
    redis_url: str = Field("redis://localhost:6379")
    database_url: str = Field("postgresql://user:pass@localhost/userstories")
    upload_dir: str = Field("./uploads")
    max_file_size: int = Field(50 * 1024 * 1024)  # 50MB

# ========================================
# 3. DATABASE MODELS
# ========================================

Base = declarative_base()

class Project(Base):
    __tablename__ = "projects"
    
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    name = Column(String, nullable=False)
    description = Column(Text)
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class Document(Base):
    __tablename__ = "documents"
    
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    project_id = Column(String, nullable=False)
    filename = Column(String, nullable=False)
    file_path = Column(String, nullable=False)
    file_hash = Column(String, nullable=False)
    file_size = Column(Integer, nullable=False)
    processed = Column(Boolean, default=False)
    created_at = Column(DateTime, default=datetime.utcnow)

class UserStory(Base):
    __tablename__ = "user_stories"
    
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    project_id = Column(String, nullable=False)
    title = Column(String, nullable=False)
    story = Column(Text, nullable=False)
    acceptance_criteria = Column(Text)
    priority = Column(String, default="medium")
    created_at = Column(DateTime, default=datetime.utcnow)

# ========================================
# 4. RAG SYSTEM IMPLEMENTATION
# ========================================

class RAGSystem:
    def __init__(self, config: AppConfig):
        self.config = config
        self.embeddings = self._init_embeddings()
        self.llm = self._init_llm()
        self.vector_store = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.rag.chunk_size,
            chunk_overlap=config.rag.chunk_overlap
        )
        self.chroma_client = chromadb.PersistentClient(path="./chroma_db")
        
    def _init_embeddings(self):
        """Initialize embeddings based on configuration"""
        llm_config = self.config.llm
        
        if llm_config.provider == "openai":
            return OpenAIEmbeddings(
                openai_api_key=llm_config.api_key,
                model="text-embedding-ada-002"
            )
        elif llm_config.provider == "azure":
            return AzureOpenAIEmbeddings(
                azure_endpoint=llm_config.base_url,
                api_key=llm_config.api_key,
                api_version=llm_config.api_version
            )
        else:
            raise ValueError(f"Unsupported embedding provider: {llm_config.provider}")
    
    def _init_llm(self):
        """Initialize LLM based on configuration"""
        llm_config = self.config.llm
        
        if llm_config.provider == "openai":
            return ChatOpenAI(
                openai_api_key=llm_config.api_key,
                model_name=llm_config.model_name,
                temperature=llm_config.temperature,
                max_tokens=llm_config.max_tokens
            )
        elif llm_config.provider == "azure":
            return AzureChatOpenAI(
                azure_endpoint=llm_config.base_url,
                api_key=llm_config.api_key,
                api_version=llm_config.api_version,
                deployment_name=llm_config.model_name,
                temperature=llm_config.temperature,
                max_tokens=llm_config.max_tokens
            )
        elif llm_config.provider == "gemini":
            return ChatGooglePalm(
                google_api_key=llm_config.api_key,
                model_name=llm_config.model_name,
                temperature=llm_config.temperature
            )
        else:
            raise ValueError(f"Unsupported LLM provider: {llm_config.provider}")
    
    async def process_document(self, file_path: str, project_id: str):
        """Process document and add to vector store"""
        try:
            # Load document based on file type
            if file_path.endswith('.pdf'):
                loader = PDFLoader(file_path)
            elif file_path.endswith('.txt'):
                loader = TextLoader(file_path)
            elif file_path.endswith(('.doc', '.docx')):
                loader = UnstructuredWordDocumentLoader(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_path}")
            
            documents = loader.load()
            
            # Split documents into chunks
            chunks = self.text_splitter.split_documents(documents)
            
            # Add metadata
            for chunk in chunks:
                chunk.metadata.update({
                    "project_id": project_id,
                    "source": file_path,
                    "processed_at": datetime.utcnow().isoformat()
                })
            
            # Create or get collection
            collection_name = f"project_{project_id}"
            collection = self.chroma_client.get_or_create_collection(
                name=collection_name,
                embedding_function=self.embeddings
            )
            
            # Add documents to vector store
            texts = [chunk.page_content for chunk in chunks]
            metadatas = [chunk.metadata for chunk in chunks]
            ids = [f"{project_id}_{i}" for i in range(len(chunks))]
            
            collection.add(
                documents=texts,
                metadatas=metadatas,
                ids=ids
            )
            
            logger.info(f"Processed {len(chunks)} chunks for project {project_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    async def query_documents(self, query: str, project_id: str) -> List[Document]:
        """Query vector store for relevant documents"""
        try:
            collection_name = f"project_{project_id}"
            collection = self.chroma_client.get_collection(name=collection_name)
            
            results = collection.query(
                query_texts=[query],
                n_results=self.config.rag.max_retrieved_docs,
                where={"project_id": project_id}
            )
            
            documents = []
            if results["documents"]:
                for doc, metadata in zip(results["documents"][0], results["metadatas"][0]):
                    documents.append(Document(
                        page_content=doc,
                        metadata=metadata
                    ))
            
            return documents
            
        except Exception as e:
            logger.error(f"Error querying documents: {str(e)}")
            return []

# ========================================
# 5. USER STORY GENERATOR
# ========================================

class UserStoryGenerator:
    def __init__(self, rag_system: RAGSystem):
        self.rag_system = rag_system
        self.prompt_template = PromptTemplate(
            input_variables=["context", "requirements", "project_context"],
            template="""
            You are an expert product manager and business analyst. Generate comprehensive user stories based on the provided requirements and context.

            Project Context: {project_context}
            
            Requirements Context from Documents:
            {context}
            
            Additional Requirements:
            {requirements}
            
            Generate user stories following this format:
            
            **Title**: [Descriptive title with emoji]
            **User Story**: As a [user type], I want [functionality] so that [benefit/value].
            **Acceptance Criteria**:
            • [Specific, testable criteria]
            • [Include edge cases and error handling]
            • [Consider security and performance requirements]
            
            **Priority**: [High/Medium/Low based on business value]
            **Epic**: [Group related stories under epics]
            **Estimation**: [Story points 1-13]
            
            Focus on:
            1. Clear, actionable user stories
            2. Comprehensive acceptance criteria
            3. Security and compliance considerations
            4. Integration points with existing systems
            5. Performance and scalability requirements
            
            Generate 5-8 detailed user stories covering the main functional areas.
            """
        )
    
    async def generate_stories(
        self, 
        project_id: str, 
        project_context: str, 
        requirements: str
    ) -> List[Dict[str, Any]]:
        """Generate user stories using RAG-enhanced context"""
        try:
            # Query relevant documents
            context_docs = await self.rag_system.query_documents(
                requirements, project_id
            )
            
            # Combine context from retrieved documents
            context = "\n\n".join([doc.page_content for doc in context_docs])
            
            # Generate user stories
            prompt = self.prompt_template.format(
                context=context,
                requirements=requirements,
                project_context=project_context
            )
            
            response = await self.rag_system.llm.agenerate([prompt])
            generated_text = response.generations[0][0].text
            
            # Parse the generated stories (simplified - you'd want more robust parsing)
            stories = self._parse_generated_stories(generated_text)
            
            return stories
            
        except Exception as e:
            logger.error(f"Error generating user stories: {str(e)}")
            raise
    
    def _parse_generated_stories(self, text: str) -> List[Dict[str, Any]]:
        """Parse generated text into structured user stories"""
        # This is a simplified parser - implement more robust parsing
        stories = []
        lines = text.split('\n')
        current_story = {}
        
        for line in lines:
            line = line.strip()
            if line.startswith('**Title**:'):
                if current_story:
                    stories.append(current_story)
                current_story = {'title': line.replace('**Title**:', '').strip()}
            elif line.startswith('**User Story**:'):
                current_story['story'] = line.replace('**User Story**:', '').strip()
            elif line.startswith('**Acceptance Criteria**:'):
                current_story['acceptance_criteria'] = []
            elif line.startswith('•') and 'acceptance_criteria' in current_story:
                current_story['acceptance_criteria'].append(line[1:].strip())
            elif line.startswith('**Priority**:'):
                current_story['priority'] = line.replace('**Priority**:', '').strip().lower()
        
        if current_story:
            stories.append(current_story)
        
        return stories

# ========================================
# 6. SECURITY IMPLEMENTATION
# ========================================

class SecurityManager:
    def __init__(self, config: SecurityConfig):
        self.config = config
        self.pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
        self.redis_client = redis.Redis.from_url(config.redis_url if hasattr(config, 'redis_url') else "redis://localhost:6379")
        
    def sanitize_input(self, text: str) -> str:
        """Sanitize user input to prevent XSS and injection attacks"""
        return bleach.clean(text, tags=[], strip=True)
    
    def validate_file_upload(self, file: UploadFile) -> bool:
        """Validate uploaded files for security"""
        allowed_extensions = {'.pdf', '.txt', '.doc', '.docx', '.md'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail="File type not allowed")
        
        if file.size > 50 * 1024 * 1024:  # 50MB limit
            raise HTTPException(status_code=400, detail="File too large")
        
        return True
    
    def check_rate_limit(self, client_ip: str) -> bool:
        """Check if client has exceeded rate limits"""
        key = f"rate_limit:{client_ip}"
        current_requests = self.redis_client.incr(key)
        
        if current_requests == 1:
            self.redis_client.expire(key, self.config.rate_limit_window)
        
        return current_requests <= self.config.rate_limit_requests
    
    def create_access_token(self, data: dict) -> str:
        """Create JWT access token"""
        to_encode = data.copy()
        expire = datetime.utcnow() + timedelta(minutes=self.config.access_token_expire_minutes)
        to_encode.update({"exp": expire})
        
        return jwt.encode(to_encode, self.config.jwt_secret, algorithm=self.config.jwt_algorithm)
    
    def verify_token(self, token: str) -> dict:
        """Verify JWT token"""
        try:
            payload = jwt.decode(token, self.config.jwt_secret, algorithms=[self.config.jwt_algorithm])
            return payload
        except jwt.ExpiredSignatureError:
            raise HTTPException(status_code=401, detail="Token expired")
        except jwt.JWTError:
            raise HTTPException(status_code=401, detail="Invalid token")

# ========================================
# 7. INTEGRATION SERVICES
# ========================================

class IntegrationService:
    def __init__(self, config: IntegrationConfig):
        self.config = config
    
    async def fetch_confluence_content(self, page_url: str) -> str:
        """Fetch content from Confluence page"""
        try:
            if not self.config.confluence_base_url:
                raise ValueError("Confluence not configured")
            
            # Extract page ID from URL
            page_id = page_url.split('/')[-1]
            
            # Confluence REST API call
            api_url = f"{self.config.confluence_base_url}/rest/api/content/{page_id}?expand=body.storage"
            
            auth = (self.config.confluence_username, self.config.confluence_api_token)
            response = requests.get(api_url, auth=auth)
            response.raise_for_status()
            
            data = response.json()
            content = data['body']['storage']['value']
            
            # Convert HTML to plain text
            soup = BeautifulSoup(content, 'html.parser')
            return soup.get_text()
            
        except Exception as e:
            logger.error(f"Error fetching Confluence content: {str(e)}")
            raise
    
    async def fetch_jira_requirements(self, project_key: str) -> List[Dict]:
        """Fetch requirements from Jira project"""
        try:
            if not self.config.jira_base_url:
                raise ValueError("Jira not configured")
            
            # Jira REST API call
            api_url = f"{self.config.jira_base_url}/rest/api/2/search"
            params = {
                'jql': f'project = {project_key} AND issuetype in (Epic, Story, Task)',
                'fields': 'summary,description,issuetype,priority,status'
            }
            
            auth = (self.config.jira_username, self.config.jira_api_token)
            response = requests.get(api_url, auth=auth, params=params)
            response.raise_for_status()
            
            data = response.json()
            issues = []
            
            for issue in data['issues']:
                issues.append({
                    'key': issue['key'],
                    'summary': issue['fields']['summary'],
                    'description': issue['fields'].get('description', ''),
                    'issue_type': issue['fields']['issuetype']['name'],
                    'priority': issue['fields']['priority']['name'],
                    'status': issue['fields']['status']['name']
                })
            
            return issues
            
        except Exception as e:
            logger.error(f"Error fetching Jira requirements: {str(e)}")
            raise

# ========================================
# 8. FASTAPI APPLICATION
# ========================================

# Load configuration
config = AppConfig(**json.loads(os.getenv('APP_CONFIG', '{}')))

# Initialize services
rag_system = RAGSystem(config)
story_generator = UserStoryGenerator(rag_system)
security_manager = SecurityManager(config.security)
integration_service = IntegrationService(config.integrations)

# Initialize FastAPI app
app = FastAPI(
    title="AI User Story Generator",
    description="Production-grade AI agent for generating user stories",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security dependency
security = HTTPBearer()

async def get_current_user(credentials: HTTPAuthorizationCredentials = Security(security)):
    """Verify JWT token and get current user"""
    payload = security_manager.verify_token(credentials.credentials)
    return payload

# Database setup
engine = create_engine(config.database_url)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base.metadata.create_all(bind=engine)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ========================================
# 9. API ENDPOINTS
# ========================================

# Request/Response models
class ProjectCreate(BaseModel):
    name: str
    description: Optional[str] = None

class UserStoryRequest(BaseModel):
    project_context: str
    requirements: str
    source_type: str = "upload"
    source_url: Optional[str] = None

class LLMConfigUpdate(BaseModel):
    provider: str
    model_name: str
    api_key: str
    base_url: Optional[str] = None
    api_version: Optional[str] = None
    temperature: float = 0.7
    max_tokens: int = 2000

@app.post("/api/projects")
async def create_project(
    project: ProjectCreate,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new project"""
    db_project = Project(
        name=project.name,
        description=project.description
    )
    db.add(db_project)
    db.commit()
    db.refresh(db_project)
    
    return {
        "id": db_project.id,
        "name": db_project.name,
        "description": db_project.description,
        "created_at": db_project.created_at
    }

@app.post("/api/projects/{project_id}/upload")
async def upload_document(
    project_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload and process a document"""
    # Validate file
    security_manager.validate_file_upload(file)
    
    # Save file
    os.makedirs(config.upload_dir, exist_ok=True)
    file_path = os.path.join(config.upload_dir, f"{project_id}_{file.filename}")
    
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # Calculate file hash
    file_hash = hashlib.sha256(content).hexdigest()
    
    # Save to database
    db_document = Document(
        project_id=project_id,
        filename=file.filename,
        file_path=file_path,
        file_hash=file_hash,
        file_size=len(content)
    )
    db.add(db_document)
    db.commit()
    
    # Process with RAG system
    try:
        await rag_system.process_document(file_path, project_id)
        db_document.processed = True
        db.commit()
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        raise HTTPException(status_code=500, detail="Document processing failed")
    
    return {
        "message": "Document uploaded and processed successfully",
        "document_id": db_document.id,
        "filename": file.filename
    }

@app.post("/api/projects/{project_id}/generate-stories")
async def generate_user_stories(
    project_id: str,
    request: UserStoryRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generate user stories for a project"""
    try:
        # Handle different source types
        requirements = request.requirements
        
        if request.source_type == "confluence" and request.source_url:
            confluence_content = await integration_service.fetch_confluence_content(request.source_url)
            requirements += f"\n\nConfluence Content:\n{confluence_content}"
        
        elif request.source_type == "jira" and request.source_url:
            jira_issues = await integration_service.fetch_jira_requirements(request.source_url)
            jira_content = "\n".join([f"{issue['summary']}: {issue['description']}" for issue in jira_issues])
            requirements += f"\n\nJira Issues:\n{jira_content}"
        
        # Generate user stories
        stories = await story_generator.generate_stories(
            project_id=project_id,
            project_context=request.project_context,
            requirements=requirements
        )
        
        # Save to database
        saved_stories = []
        for story in stories:
            db_story = UserStory(
                project_id=project_id,
                title=story.get('title', ''),
                story=story.get('story', ''),
                acceptance_criteria=json.dumps(story.get('acceptance_criteria', [])),
                priority=story.get('priority', 'medium')
            )
            db.add(db_story)
            saved_stories.append(db_story)
        
        db.commit()
        
        return {
            "message": f"Generated {len(stories)} user stories",
            "stories": [
                {
                    "id": story.id,
                    "title": story.title,
                    "story": story.story,
                    "acceptance_criteria": json.loads(story.acceptance_criteria),
                    "priority": story.priority
                }
                for story in saved_stories
            ]
        }
        
    except Exception as e:
        logger.error(f"Error generating user stories: {str(e)}")
        raise HTTPException(status_code=500, detail="Story generation failed")

@app.put("/api/config/llm")
async def update_llm_config(
    config_update: LLMConfigUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Update LLM configuration"""
    try:
        # Update global config
        config.llm.provider = config_update.provider
        config.llm.model_name = config_update.model_name
        config.llm.api_key = config_update.api_key
        config.llm.base_url = config_update.base_url
        config.llm.api_version = config_update.api_version
        config.llm.temperature = config_update.temperature
        config.llm.max_tokens = config_update.max_tokens
        
        # Reinitialize RAG system with new config
        global rag_system, story_generator
        rag_system = RAGSystem(config)
        story_generator = UserStoryGenerator(rag_system)
        
        return {"message": "LLM configuration updated successfully"}
        
    except Exception as e:
        logger.error(f"Error updating LLM config: {str(e)}")
        raise HTTPException(status_code=500, detail="Configuration update failed")

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "1.0.0"
    }

@app.get("/metrics")
async def metrics():
    """Prometheus metrics endpoint"""
    return generate_latest()

# Serve static files
app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)